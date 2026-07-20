from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any, Literal

from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, Response, UploadFile, status
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.database import get_db
from app.schemas.studio import (
    ArtifactDecision,
    ArtifactUpdate,
    ChapterTreeRepairRequest,
    ChatRequest,
    ContinuationImportRequest,
    ContinuationSettingsUpdate,
    GenerateRequest,
    MessageProposalDecision,
    OutlineImportRequest,
    ProviderSecretUpdate,
    ProviderSetup,
    SnapshotCreate,
    StudioProjectCreate,
    StudioStateUpdate,
)
from app.services import studio


router = APIRouter(prefix="/studio", tags=["studio-v2"])
MAX_OUTLINE_BYTES = 10 * 1024 * 1024


@router.get("/projects")
def list_projects(db: Session = Depends(get_db)) -> list[dict[str, Any]]:
    return studio.dashboard(db)


@router.post("/projects", status_code=status.HTTP_201_CREATED)
def create_project(
    payload: StudioProjectCreate, db: Session = Depends(get_db)
) -> dict[str, Any]:
    with db.begin():
        return studio.create_project(db, payload)


@router.post("/continuations", status_code=status.HTTP_201_CREATED)
def create_continuation(
    payload: ContinuationImportRequest, db: Session = Depends(get_db)
) -> dict[str, Any]:
    with db.begin():
        return studio.create_continuation_project(db, payload)


@router.post("/continuations/file", status_code=status.HTTP_201_CREATED)
async def create_continuation_from_file(
    title: str = Form(...),
    file: UploadFile = File(...),
    target_words: int | None = Form(None),
    target_chapters: int | None = Form(None),
    target_volumes: int | None = Form(None),
    continuation_start: Literal["choose", "current", "next"] = Form("choose"),
    direction_mode: Literal["user", "ai", "switchable"] = Form("switchable"),
    user_outline: str = Form(""),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    content = await file.read(MAX_OUTLINE_BYTES + 1)
    if len(content) > MAX_OUTLINE_BYTES:
        raise HTTPException(status_code=413, detail="半成品小说文件不得超过 10 MB")
    text = _extract_document_text(content, file.filename or "导入正文")
    payload = ContinuationImportRequest(
        title=title,
        text=text,
        source_name=file.filename or "导入正文",
        target_words=target_words,
        target_chapters=target_chapters,
        target_volumes=target_volumes,
        continuation_start=continuation_start,
        direction_mode=direction_mode,
        user_outline=user_outline,
    )
    with db.begin():
        return studio.create_continuation_project(db, payload)


@router.get("/projects/{project_id}")
def read_project(project_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    return studio.project_overview(db, project_id)


@router.patch("/projects/{project_id}/state")
def update_project_state(
    project_id: int,
    payload: StudioStateUpdate,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    with db.begin():
        return studio.update_state(db, project_id, payload)


@router.patch("/projects/{project_id}/continuation/settings")
def update_continuation_settings(
    project_id: int,
    payload: ContinuationSettingsUpdate,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    with db.begin():
        return studio.update_continuation_settings(db, project_id, payload)


@router.delete("/projects/{project_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_project(project_id: int, db: Session = Depends(get_db)) -> Response:
    project = studio._project(db, project_id)
    with db.begin():
        from datetime import datetime, timezone

        project.deleted_at = datetime.now(timezone.utc)
        project.revision += 1
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@router.put("/artifacts/{artifact_id}")
def update_artifact(
    artifact_id: int,
    payload: ArtifactUpdate,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    with db.begin():
        return studio.update_artifact(db, artifact_id, payload)


@router.get("/artifacts/{artifact_id}/versions")
def artifact_versions(artifact_id: int, db: Session = Depends(get_db)) -> list[dict[str, Any]]:
    return studio.artifact_versions(db, artifact_id)


@router.post("/artifacts/{artifact_id}/decision")
def decide_artifact(
    artifact_id: int,
    payload: ArtifactDecision,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    with db.begin():
        return studio.decide_artifact(db, artifact_id, payload)


@router.post("/projects/{project_id}/generate/{phase}")
async def generate_phase(
    project_id: int,
    phase: str,
    payload: GenerateRequest,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    return await studio.generate(db, project_id, phase, payload)


@router.post("/projects/{project_id}/chat")
async def chat(
    project_id: int,
    payload: ChatRequest,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    return await studio.chat(db, project_id, payload)


@router.post("/projects/{project_id}/messages/{message_id}/proposal")
def decide_message_proposal(
    project_id: int,
    message_id: int,
    payload: MessageProposalDecision,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    with db.begin():
        return studio.decide_message_proposal(db, project_id, message_id, payload.action)


@router.post("/projects/{project_id}/outline/preview")
def preview_outline(
    project_id: int,
    payload: OutlineImportRequest,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    project = studio._project(db, project_id)
    return studio.parse_outline(payload.text, project.title)


@router.post("/projects/{project_id}/outline/preview-file")
async def preview_outline_file(
    project_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    project = studio._project(db, project_id)
    content = await file.read(MAX_OUTLINE_BYTES + 1)
    if len(content) > MAX_OUTLINE_BYTES:
        raise HTTPException(status_code=413, detail="大纲文件不得超过 10 MB")
    suffix = Path(file.filename or "").suffix.lower()
    try:
        if suffix == ".docx":
            document = Document(BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        elif suffix in {".txt", ".md", ".markdown"}:
            text = content.decode("utf-8-sig")
        else:
            raise HTTPException(status_code=415, detail="仅支持 TXT、Markdown 和 DOCX 文件")
    except UnicodeDecodeError as exc:
        raise HTTPException(status_code=422, detail="文本文件必须使用 UTF-8 编码") from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail="Word 文件损坏或格式不受支持") from exc
    return {**studio.parse_outline(text, project.title), "source_text": text}


@router.post("/projects/{project_id}/style-reference")
async def upload_style_reference(
    project_id: int,
    file: UploadFile = File(...),
    use_demo_model: bool = Form(False),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    studio._project(db, project_id)
    content = await file.read(MAX_OUTLINE_BYTES + 1)
    if len(content) > MAX_OUTLINE_BYTES:
        raise HTTPException(status_code=413, detail="参考文风文件不得超过 10 MB")
    suffix = Path(file.filename or "").suffix.lower()
    try:
        if suffix == ".docx":
            document = Document(BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        elif suffix in {".txt", ".md", ".markdown"}:
            text = content.decode("utf-8-sig")
        else:
            raise HTTPException(status_code=415, detail="参考文风仅支持 TXT、Markdown 和 DOCX")
    except UnicodeDecodeError as exc:
        raise HTTPException(status_code=422, detail="文本文件必须使用 UTF-8 编码") from exc
    if not text.strip():
        raise HTTPException(status_code=422, detail="参考文风文件没有可分析的文本")
    return await studio.extract_style_reference(
        db, project_id, text[:200_000], file.filename or "参考文本", use_demo_model
    )


@router.post("/projects/{project_id}/outline/import")
def import_outline(
    project_id: int,
    payload: OutlineImportRequest,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    with db.begin():
        return studio.import_outline(db, project_id, payload)


@router.post("/projects/{project_id}/snapshots", status_code=status.HTTP_201_CREATED)
def create_snapshot(
    project_id: int,
    payload: SnapshotCreate,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    with db.begin():
        return studio.create_snapshot(db, project_id, payload)


@router.post("/projects/{project_id}/snapshots/{snapshot_id}/restore")
def restore_snapshot(
    project_id: int,
    snapshot_id: int,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    with db.begin():
        return studio.restore_snapshot(db, project_id, snapshot_id)


@router.get("/projects/{project_id}/chapter-tree/repair-preview")
def chapter_tree_repair_preview(
    project_id: int, db: Session = Depends(get_db)
) -> dict[str, Any]:
    return studio.chapter_tree_repair_preview(db, project_id)


@router.post("/projects/{project_id}/chapter-tree/repair")
def repair_chapter_tree(
    project_id: int,
    payload: ChapterTreeRepairRequest,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    with db.begin():
        return studio.repair_chapter_tree(db, project_id, payload)


@router.get("/providers")
def list_providers(db: Session = Depends(get_db)) -> list[dict[str, Any]]:
    return studio.list_studio_providers(db)


@router.post("/providers", status_code=status.HTTP_201_CREATED)
def setup_provider(
    payload: ProviderSetup, db: Session = Depends(get_db)
) -> dict[str, Any]:
    with db.begin():
        return studio.setup_provider(db, payload)


@router.put("/providers/{provider_id}/secret")
def update_provider_secret(
    provider_id: int,
    payload: ProviderSecretUpdate,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    return studio.update_provider_secret(db, provider_id, payload.api_key)


@router.delete("/providers/{provider_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_provider(provider_id: int, db: Session = Depends(get_db)) -> Response:
    with db.begin():
        studio.delete_studio_provider(db, provider_id)
    return Response(status_code=status.HTTP_204_NO_CONTENT)


def _extract_document_text(content: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    try:
        if suffix == ".docx":
            document = Document(BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        elif suffix == ".pdf":
            reader = PdfReader(BytesIO(content))
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        elif suffix in {".txt", ".md", ".markdown"}:
            text = content.decode("utf-8-sig")
        else:
            raise HTTPException(status_code=415, detail="仅支持 TXT、Markdown、DOCX 和 PDF 文件")
    except UnicodeDecodeError as exc:
        raise HTTPException(status_code=422, detail="文本文件必须使用 UTF-8 编码") from exc
    except (ValueError, OSError) as exc:
        raise HTTPException(status_code=422, detail="文件损坏、加密或格式不受支持") from exc
    if not text.strip():
        raise HTTPException(status_code=422, detail="文件中没有可识别的正文文本")
    return text.strip()
